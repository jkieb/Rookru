"""Tests, die ohne Netzwerk, API-Schlüssel und private Dateien laufen.

Die docx-Vorlagen werden im Test selbst erzeugt, damit nichts aus privat/
gebraucht wird.
"""

from __future__ import annotations

import json
from pathlib import Path

import pytest
from docx import Document
from docx.oxml.ns import qn

from rookru import screening
from rookru.compose import ComposerError, StubComposer, detect_focus, message_text, parse_response
from rookru.config import (
    ConfigError,
    FocusRule,
    LetterSettings,
    SearchSettings,
    load_settings,
    slugify,
)
from rookru.models import CVAdaptation, Job, LetterContent, Screening, SectionEdit, TemplateData
from rookru import pipeline
from rookru.pipeline import (
    build_application,
    compact_name,
    describe_address,
    describe_fill,
    describe_fit,
    save_search_run,
)
from rookru.render import cv as cv_render
from rookru.render import docx_tools as dt
from rookru.render import letter as letter_render
from rookru.render.bundle import BundleError, merge_pdfs
from rookru.sources import anfragen, rank_jobs, search_all
from rookru.sources import careerjet as cj
from rookru.sources import eures as eu
from rookru.sources import jooble as jb
from rookru.sources.common import UNBEKANNT, clean_text, zu_alt
from rookru.sources.adzuna import build_url
from rookru.sources.common import dedup_key
from rookru.sources.local import load_jobs_file

# ---------------------------------------------------------------- Hilfsvorlagen


def make_letter_template(path: Path) -> Path:
    document = Document()
    for text in (
        "Max Muster",
        "Wien, {{DATUM}}",
        "{{FIRMA_NAME}}",
        "{{FIRMA_ABTEILUNG}}",
        "{{FIRMA_STRASSE}}",
        "{{FIRMA_PLZ_ORT}}",
        "Bewerbung als {{STELLE}}",
        "Sehr geehrte Damen und Herren,",
        "{{ABSATZ_1}}",
        "{{ABSATZ_2}}",
        "Mit freundlichen Grüßen,",
    ):
        document.add_paragraph(text)
    document.save(str(path))
    return path


def make_cv_template(path: Path) -> Path:
    document = Document()
    document.add_paragraph("Max Muster")
    document.add_paragraph("AUSBILDUNG")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "2023 – dato"
    cell = table.rows[0].cells[1]
    cell.text = "Bachelorstudium Maschinenbau, TU Wien"
    cell.add_paragraph("●  Bachelorarbeit zu additiver Fertigung")
    document.add_paragraph("")
    document.add_paragraph("PROJEKTE")
    for anchor, head, bullets in (
        ("Drucker", "FDM-3D-Druck, privat", ["●  Halterungen konstruiert", "●  PETG und TPU"]),
        ("GitHub", "github.com/muster", ["●  Python-Projekte"]),
    ):
        table = document.add_table(rows=1, cols=2)
        table.rows[0].cells[0].text = anchor
        cell = table.rows[0].cells[1]
        cell.text = head
        for bullet in bullets:
            cell.add_paragraph(bullet)
        document.add_paragraph("")
    document.add_paragraph("SPRACHKENNTNISSE")
    document.add_paragraph("Deutsch: Muttersprache")
    document.add_paragraph("BESONDERE KENNTNISSE UND FÄHIGKEITEN")
    document.add_paragraph("CAD-Kenntnisse: Creo")
    document.add_paragraph("Programmierkenntnisse: Python")
    document.save(str(path))
    return path


@pytest.fixture
def cv_template(tmp_path: Path) -> Path:
    return make_cv_template(tmp_path / "lebenslauf.docx")


@pytest.fixture
def letter_template(tmp_path: Path) -> Path:
    return make_letter_template(tmp_path / "brief.docx")


# ------------------------------------------------------------------ docx_tools


def test_platzhalter_ueber_mehrere_runs(tmp_path: Path) -> None:
    """Word zerlegt Platzhalter oft in mehrere Runs — das muss trotzdem greifen."""
    document = Document()
    paragraph = document.add_paragraph()
    for teil in ("{{FIR", "MA_", "NAME}}"):
        paragraph.add_run(teil)
    dt.replace_in_paragraph(paragraph, {"{{FIRMA_NAME}}": "ACME GmbH"})
    assert paragraph.text == "ACME GmbH"


def test_platzhalter_behaelt_formatierung(tmp_path: Path) -> None:
    document = Document()
    paragraph = document.add_paragraph()
    run = paragraph.add_run("{{STELLE}}")
    run.bold = True
    dt.replace_in_paragraph(paragraph, {"{{STELLE}}": "Werkstudent"})
    assert paragraph.runs[0].bold is True
    assert paragraph.text == "Werkstudent"


def test_umgebender_text_bleibt_erhalten() -> None:
    document = Document()
    paragraph = document.add_paragraph("Bewerbung als {{STELLE}} (m/w/d)")
    dt.replace_in_paragraph(paragraph, {"{{STELLE}}": "Werkstudent"})
    assert paragraph.text == "Bewerbung als Werkstudent (m/w/d)"


def test_bullets_ersetzen_und_ergaenzen(cv_template: Path) -> None:
    document = Document(str(cv_template))
    tables = dt.section_tables(document, "PROJEKTE", ["SPRACHKENNTNISSE"])
    cell = tables[0].rows[0].cells[1]
    dt.set_bullets(cell, ["Erster Punkt", "Zweiter Punkt", "Dritter Punkt"])
    bullets = [p.text for p in dt.cell_bullets(cell)]
    assert bullets == ["●  Erster Punkt", "●  Zweiter Punkt", "●  Dritter Punkt"]


def test_bullets_kuerzen(cv_template: Path) -> None:
    document = Document(str(cv_template))
    tables = dt.section_tables(document, "PROJEKTE", ["SPRACHKENNTNISSE"])
    cell = tables[0].rows[0].cells[1]
    dt.set_bullets(cell, ["Nur einer"])
    assert len(dt.cell_bullets(cell)) == 1


# ------------------------------------------------------------------- Lebenslauf


def test_projekte_umsortieren(cv_template: Path, tmp_path: Path) -> None:
    adaptation = CVAdaptation(project_order=["GitHub", "Drucker"])
    output, warnings = cv_render.render_cv(cv_template, adaptation, tmp_path / "out.docx")
    assert warnings == []
    assert cv_render.project_anchors(output) == ["GitHub", "Drucker"]


def test_unbekanntes_projekt_warnt_statt_zu_scheitern(
    cv_template: Path, tmp_path: Path
) -> None:
    adaptation = CVAdaptation(project_edits=[SectionEdit(anchor="Gibt es nicht", bullets=["x"])])
    _, warnings = cv_render.render_cv(cv_template, adaptation, tmp_path / "out.docx")
    assert any("nicht in der Vorlage" in w for w in warnings)


def test_unveraenderte_abschnitte_bleiben_gleich(cv_template: Path, tmp_path: Path) -> None:
    """Sprachkenntnisse dürfen sich nie ändern — sie sind nicht freigegeben."""
    adaptation = CVAdaptation(project_order=["GitHub", "Drucker"], skill_lines=["CAD: Creo"])
    output, _ = cv_render.render_cv(cv_template, adaptation, tmp_path / "out.docx")
    texte = [p.text for p in Document(str(output)).paragraphs]
    assert "Deutsch: Muttersprache" in texte


def test_kenntnis_zeilen_ersetzen(cv_template: Path, tmp_path: Path) -> None:
    adaptation = CVAdaptation(skill_lines=["3D-Druck: Prusa", "CAD: Fusion 360", "Extra: Zeile"])
    output, _ = cv_render.render_cv(cv_template, adaptation, tmp_path / "out.docx")
    assert cv_render.skill_lines(output) == [
        "3D-Druck: Prusa",
        "CAD: Fusion 360",
        "Extra: Zeile",
    ]


def test_faktenblatt_enthaelt_alle_abschnitte(cv_template: Path) -> None:
    facts = cv_render.template_facts(cv_template)
    assert "## PROJEKTE" in facts
    assert "Halterungen konstruiert" in facts


# -------------------------------------------------------- Motivationsschreiben


def test_brief_fuellt_alle_platzhalter(letter_template: Path, tmp_path: Path) -> None:
    content = LetterContent(
        subject="Werkstudent Maschinenbau",
        salutation="Sehr geehrte Frau Muster,",
        paragraphs=["Erster Absatz.", "Zweiter Absatz."],
        company_name="ACME GmbH",
        company_department="Personal",
        company_street="Teststraße 1",
        company_postal_city="1010 Wien",
    )
    output = letter_render.render_letter(letter_template, content, tmp_path / "brief_out.docx")
    assert letter_render.check_placeholders(output) == set()
    text = "\n".join(p.text for p in Document(str(output)).paragraphs)
    assert "Bewerbung als Werkstudent Maschinenbau" in text
    assert "ACME GmbH" in text


def test_mehr_absaetze_als_platzhalter(letter_template: Path, tmp_path: Path) -> None:
    """Die Vorlage hat zwei Absatz-Slots; drei Absätze müssen trotzdem passen."""
    content = LetterContent(
        subject="Stelle",
        salutation="Sehr geehrte Damen und Herren,",
        paragraphs=["Eins.", "Zwei.", "Drei."],
        company_name="ACME GmbH",
    )
    output = letter_render.render_letter(letter_template, content, tmp_path / "brief3.docx")
    text = "\n".join(p.text for p in Document(str(output)).paragraphs)
    assert letter_render.check_placeholders(output) == set()
    for absatz in ("Eins.", "Zwei.", "Drei."):
        assert absatz in text


def test_leere_adresszeile_wird_entfernt(letter_template: Path, tmp_path: Path) -> None:
    content = LetterContent(
        subject="Stelle",
        salutation="Sehr geehrte Damen und Herren,",
        paragraphs=["Text."],
        company_name="ACME GmbH",
    )
    output = letter_render.render_letter(letter_template, content, tmp_path / "brief_kurz.docx")
    texte = [p.text for p in Document(str(output)).paragraphs]
    assert "" not in [t for t in texte if "{{" in t]
    assert letter_render.check_placeholders(output) == set()


def test_datum_deutsch_formatiert() -> None:
    from datetime import date

    assert letter_render.format_date(date(2026, 8, 13)) == "13. August 2026"


# ------------------------------------------------------------------ Schwerpunkt


def _rules() -> list[FocusRule]:
    return [
        FocusRule("it", "IT", ["python", "sql"], ["GitHub"]),
        FocusRule("konstruktion", "Konstruktion", ["cad", "3d-druck"], ["Prusa MK3S+"]),
    ]


def test_schwerpunkt_it() -> None:
    job = Job(id="1", title="Werkstudent Datenauswertung", company="X",
              description="Python und SQL für Messdaten")
    assert detect_focus(job, _rules()).key == "it"


def test_schwerpunkt_konstruktion() -> None:
    job = Job(id="2", title="Aushilfe 3D-Druck", company="Y",
              description="CAD-Konstruktion von Bauteilen")
    assert detect_focus(job, _rules()).key == "konstruktion"


def test_ohne_treffer_kein_schwerpunkt() -> None:
    job = Job(id="3", title="Bürokraft", company="Z", description="Ablage und Telefon")
    assert detect_focus(job, _rules()) is None


# ----------------------------------------------------------------- Bündel & Co.


def test_buendel_meldet_fehlende_anlage(tmp_path: Path) -> None:
    with pytest.raises(BundleError, match="Anlage fehlt"):
        merge_pdfs([("Zeugnis", tmp_path / "gibtsnicht.pdf")], tmp_path / "b.pdf")


def test_buendel_ohne_teile() -> None:
    with pytest.raises(BundleError):
        merge_pdfs([], Path("/tmp/leer.pdf"))


def test_dateinamen_kompakt() -> None:
    assert compact_name("Knorr-Bremse GmbH") == "KnorrBremse"
    assert compact_name("Beispiel Automation GmbH") == "BeispielAutomation"


def test_slugify() -> None:
    assert slugify("Werkstudent Maschinenbau (m/w/d)") == "werkstudent-maschinenbau-m-w-d"


def test_fit_meldung() -> None:
    assert describe_fit("Lebenslauf", (1.0, 1.0)) == []
    nur_abstaende = describe_fit("Lebenslauf", (1.0, 0.6))[0]
    assert "60%" in nur_abstaende and "Schrift" not in nur_abstaende
    auch_schrift = describe_fit("Lebenslauf", (0.9, 0.5))[0]
    assert "Schrift" in auch_schrift and "90%" in auch_schrift


def test_mehrwortsuche_verlangt_alle_woerter() -> None:
    url = build_url(SearchSettings(), "Werkstudent Maschinenbau", 1, "id", "key")
    assert "what_and=Werkstudent+Maschinenbau" in url
    assert "what=" not in url


def test_einwortsuche_bleibt_bei_what() -> None:
    url = build_url(SearchSettings(), "Maschinenbau", 1, "id", "key")
    assert "what=Maschinenbau" in url


def test_suchbegriff_im_titel_schlaegt_schwerpunkt() -> None:
    """Die gesuchte Werkstudentenstelle gehört vor eine Vollzeitstelle mit mehr Fachbegriffen."""
    werkstudent = Job(id="1", title="Werkstudent Maschinenbau (m/w/d)", company="A")
    vollzeit = Job(
        id="2",
        title="Leitung Konstruktion",
        company="B",
        description="cad creo konstruktion fertigung entwicklung",
    )
    ranked = rank_jobs([vollzeit, werkstudent], _rules(), queries=["Werkstudent Maschinenbau"])
    assert [job.id for job, _, _ in ranked] == ["1", "2"]
    assert ranked[0][2] == 2  # beide Suchwörter im Titel


def test_ranking_ohne_suchbegriffe_bleibt_beim_schwerpunkt() -> None:
    treffer = Job(id="1", title="Konstrukteur", company="A", description="cad creo 3d-druck")
    ohne = Job(id="2", title="Bürokraft", company="B")
    ranked = rank_jobs([ohne, treffer], _rules())
    assert [job.id for job, _, _ in ranked] == ["1", "2"]


def test_wortspanne_leitet_untergrenze_ab() -> None:
    assert LetterSettings(max_words=300).target_words() == (255, 300)
    assert LetterSettings(max_words=330, min_words=280).target_words() == (280, 330)


def test_untergrenze_ueber_obergrenze_wird_gedeckelt() -> None:
    assert LetterSettings(max_words=200, min_words=400).target_words() == (200, 200)


def test_gestauchter_brief_bekommt_keinen_zu_kurz_rat(
    letter_template: Path, cv_template: Path, tmp_path: Path, monkeypatch
) -> None:
    """Wurde gestaucht, war der Brief zu lang — 'min_woerter anheben' wäre verkehrt."""
    settings = _settings_for_parse(tmp_path, cv_template, letter_template)
    job = Job(id="1", title="Werkstudent", company="ACME")

    monkeypatch.setattr(
        pipeline.convert, "fit_to_one_page", lambda *a, **k: (tmp_path / "x.pdf", 1, (1.0, 0.70))
    )
    monkeypatch.setattr(pipeline.convert, "fill_ratio", lambda *a, **k: 0.60)
    monkeypatch.setattr(pipeline.convert, "page_count", lambda *a, **k: 1)
    monkeypatch.setattr(pipeline.bundle, "merge_pdfs", lambda *a, **k: None)

    app = build_application(settings, job, StubComposer(), output_root=tmp_path / "out")
    assert not any("min_woerter" in w for w in app.warnings)
    assert any("gestaucht" in w for w in app.warnings)


def test_halb_leere_seite_warnt() -> None:
    assert describe_fill(0.99, 300, 330) == []
    warnung = describe_fill(0.62, 190, 330)[0]
    assert "62%" in warnung and "min_woerter" in warnung


# ---------------------------------------------------------------- Careerjet


def test_careerjet_url_enthaelt_ort_und_umkreis() -> None:
    settings = SearchSettings(country="at", where="Wien", distance_km=25, results=20)
    url = cj.build_url(settings, "Werkstudent Maschinenbau", 1)
    assert "locale_code=de_AT" in url
    assert "location=Wien" in url and "radius=25" in url
    assert "keywords=Werkstudent+Maschinenbau" in url


def test_careerjet_locale_aus_land_ableitbar_und_ueberschreibbar() -> None:
    assert cj.locale_for(SearchSettings(country="de")) == "de_DE"
    assert cj.locale_for(SearchSettings(country="fr")) == "en_GB"
    assert cj.locale_for(SearchSettings(country="at", locale="en_GB")) == "en_GB"


def test_careerjet_treffer_wird_zu_job() -> None:
    job = cj._to_job({
        "title": "<b>Werkstudent</b>:in Konstruktion",
        "company": "ACME &amp; Co",
        "locations": "Wien, Wien",
        "description": "Laufendes Studium <b>Maschinenbau</b>",
        "date": "Wed, 05 Aug 2026 07:25:37 GMT",
        "url": "https://jobviewtrack.com/v2/abcdef",
    })
    assert job.title == "Werkstudent:in Konstruktion"  # <b> entfernt
    assert job.company == "ACME & Co"
    assert job.created == "2026-08-05"
    assert job.source == "careerjet"
    assert job.id


def test_altersfilter_fuer_quellen_ohne_eigenen() -> None:
    from datetime import date, timedelta

    frisch = Job(id="1", title="T", company="C", created=date.today().isoformat())
    alt = Job(id="2", title="T", company="C", created=(date.today() - timedelta(days=90)).isoformat())
    assert zu_alt(frisch, 30) is False
    assert zu_alt(alt, 30) is True
    assert zu_alt(alt, 0) is False  # ohne Filter bleibt alles
    assert zu_alt(Job(id="3", title="T", company="C"), 30) is False  # ohne Datum nichts wegwerfen


def test_markup_wird_aus_anzeigentexten_entfernt() -> None:
    assert clean_text("<b>Werkstudent</b>:in") == "Werkstudent:in"
    assert clean_text("ACME &amp; Co") == "ACME & Co"
    assert clean_text("mehrere    Leerzeichen\nund Umbrüche") == "mehrere Leerzeichen und Umbrüche"


# ------------------------------------------------------------------- Jooble


def test_jooble_body_enthaelt_ort_und_erlaubten_umkreis() -> None:
    settings = SearchSettings(where="Wien", distance_km=25, results=20)
    body = jb.build_body(settings, "Werkstudent Maschinenbau", 1)
    assert body["keywords"] == "Werkstudent Maschinenbau"
    assert body["location"] == "Wien"
    assert body["radius"] == "26"  # 25 ist keine erlaubte Stufe
    assert body["ResultOnPage"] == 20


def test_jooble_umkreis_wird_auf_erlaubte_stufe_gerundet() -> None:
    assert jb.radius_stufe(0) == 0
    assert jb.radius_stufe(25) == 26
    assert jb.radius_stufe(10) == 8
    assert jb.radius_stufe(500) == 80  # größer als die höchste Stufe


def test_jooble_ohne_ort_kein_umkreis() -> None:
    assert "radius" not in jb.build_body(SearchSettings(distance_km=25), "Maschinenbau", 1)


def test_jooble_treffer_wird_zu_job() -> None:
    job = jb._to_job({
        "id": "-1234567890",
        "title": "<b>Werkstudent</b> Maschinenbau (m/w/d)",
        "company": "ACME &amp; Co",
        "location": "Wien",
        "snippet": "Laufendes Studium <b>Maschinenbau</b>...",
        "link": "https://at.jooble.org/jdp/-1234567890",
        "updated": "2026-08-14T12:55:35.3870000",
    })
    assert job.title == "Werkstudent Maschinenbau (m/w/d)"
    assert job.company == "ACME & Co"
    assert job.created == "2026-08-14"  # siebenstellige Bruchsekunden abgeschnitten
    assert job.source == "jooble"


def test_jooble_unbrauchbares_datum_bleibt_leer() -> None:
    assert jb._date("") == ""
    assert jb._date("gestern") == ""


def test_user_agent_ist_ascii() -> None:
    """Ein Umlaut in der Kopfzeile lässt Jooble die Anfrage mit HTTP 400 abweisen."""
    from rookru.sources.common import USER_AGENT

    USER_AGENT.encode("ascii")  # wirft bei Nicht-ASCII


def test_jooble_ohne_schluessel_meldet_klar(monkeypatch) -> None:
    monkeypatch.delenv("JOOBLE_API_KEY", raising=False)
    with pytest.raises(jb.JoobleError, match="jooble.org/api/about"):
        jb.api_key()


def test_dieselbe_stelle_aus_zwei_quellen_zaehlt_einmal() -> None:
    a = Job(id="ad-1", title="Werkstudent Maschinenbau", company="ACME GmbH", source="adzuna")
    b = Job(id="cj-9", title="Werkstudent Maschinenbau", company="ACME GmbH", source="careerjet")
    assert dedup_key(a) == dedup_key(b)


# -------------------------------------------------------------------- EURES


def test_eures_bindet_die_rolle_an_den_titel() -> None:
    """Ohne Titelbindung matchen kurze Themen wie 'CAD' quer durch den Bestand."""
    body = eu.build_body(SearchSettings(where="Wien"), "Praktikum CAD", 1)
    assert body["keywords"] == [
        {"keyword": "Praktikum", "specificSearchCode": "TITLE"},
        {"keyword": "CAD", "specificSearchCode": "EVERYWHERE"},
    ]


def test_eures_ort_wird_zu_nuts_code() -> None:
    assert eu.location_codes(SearchSettings(where="Wien")) == ["AT13"]
    assert eu.location_codes(SearchSettings(where="  GRAZ ")) == ["AT22"]
    # Unbekannter Ort: landesweit statt gar nicht
    assert eu.location_codes(SearchSettings(where="Hintertupfing", country="at")) == ["at"]


def test_eures_zeitraum_aus_max_tagen() -> None:
    assert eu.publication_period(1) == "LAST_DAY"
    assert eu.publication_period(30) == "LAST_MONTH"
    assert eu.publication_period(90) is None  # kein passender Zeitraum → kein Filter
    assert eu.publication_period(0) is None


def test_eures_millisekunden_werden_zum_datum() -> None:
    assert eu._date(1785786302591) == "2026-08-03"
    assert eu._date(None) == ""
    assert eu._date("kaputt") == ""


def test_eures_platzhalter_firma_wird_kenntlich() -> None:
    assert eu._company({"employer": {"name": "siehe Beschreibung"}}) == UNBEKANNT
    assert eu._company({}) == UNBEKANNT
    assert eu._company({"employer": {"name": "Siemens Energy"}}) == "Siemens Energy"


def test_stellen_ohne_firmennamen_bleiben_unterscheidbar() -> None:
    """Sonst verschmelzen zwei verschiedene Anzeigen mit gleichem Titel."""
    a = Job(id="1", title="Praktikum Technik", company=UNBEKANNT)
    b = Job(id="2", title="Praktikum Technik", company=UNBEKANNT)
    assert dedup_key(a) != dedup_key(b)


def test_unbekannter_arbeitgeber_wird_gemeldet() -> None:
    warnungen = describe_address(Job(id="1", title="T", company=UNBEKANNT))
    assert any("Arbeitgeber nicht genannt" in w for w in warnungen)


def test_anfragen_zaehlt_begriffe_mal_quellen() -> None:
    einstellungen = SearchSettings(
        sources=["adzuna", "jooble", "gibtsnicht"], queries=["a", "b", "c", "d"]
    )
    assert anfragen(einstellungen) == 8  # unbekannte Quelle zählt nicht mit


def test_fortschritt_meldet_jede_abfrage(monkeypatch) -> None:
    from rookru import sources

    def quelle(settings, melden=None):
        for query in settings.queries:
            if melden:
                melden(query)
        return []

    monkeypatch.setitem(sources.QUELLEN, "adzuna", quelle)
    monkeypatch.setitem(sources.QUELLEN, "jooble", quelle)
    einstellungen = SearchSettings(sources=["adzuna", "jooble"], queries=["a", "b", "c"])

    gemeldet: list[tuple[str, str]] = []
    search_all(einstellungen, melden=lambda q, t: gemeldet.append((q, t)))
    assert len(gemeldet) == anfragen(einstellungen) == 6
    assert gemeldet[0] == ("adzuna", "a") and gemeldet[-1] == ("jooble", "c")


def test_ausgefallene_quelle_laesst_den_balken_nicht_haengen(monkeypatch) -> None:
    """Bricht eine Quelle ab, müssen ihre restlichen Schritte trotzdem zählen."""
    from rookru import sources

    def kaputt_nach_einem(settings, melden=None):
        if melden:
            melden(settings.queries[0])
        raise sources.JoobleError("Schlüssel abgelehnt")

    monkeypatch.setitem(sources.QUELLEN, "jooble", kaputt_nach_einem)
    monkeypatch.setitem(
        sources.QUELLEN, "adzuna", lambda s, melden=None: [Job(id="1", title="T", company="C")]
    )
    einstellungen = SearchSettings(sources=["adzuna", "jooble"], queries=["a", "b", "c"])

    gemeldet: list[tuple[str, str]] = []
    search_all(einstellungen, melden=lambda q, t: gemeldet.append((q, t)))
    jooble_schritte = [g for g in gemeldet if g[0] == "jooble"]
    assert len(jooble_schritte) == 3  # einer echt, zwei nachgetragen


def test_ausfall_einer_quelle_verwirft_die_andere_nicht(monkeypatch) -> None:
    from rookru import sources

    treffer = [Job(id="1", title="Werkstudent Maschinenbau", company="ACME")]
    monkeypatch.setitem(sources.QUELLEN, "adzuna", lambda s, melden=None: treffer)

    def kaputt(_settings, melden=None):
        raise sources.CareerjetError("API-Key prüfen")

    monkeypatch.setitem(sources.QUELLEN, "careerjet", kaputt)
    jobs, probleme = search_all(SearchSettings(sources=["adzuna", "careerjet"]))
    assert [j.id for j in jobs] == ["1"]
    assert any("API-Key" in p for p in probleme)


def test_alle_quellen_kaputt_ist_ein_fehler(monkeypatch) -> None:
    from rookru import sources

    def kaputt(_settings, melden=None):
        raise sources.AdzunaError("nicht erreichbar")

    monkeypatch.setitem(sources.QUELLEN, "adzuna", kaputt)
    with pytest.raises(sources.SourceError, match="nicht erreichbar"):
        search_all(SearchSettings(sources=["adzuna"]))


def test_unbekannte_quelle_meldet_sich() -> None:
    with pytest.raises(Exception, match="Unbekannte Quelle"):
        search_all(SearchSettings(sources=["monster"]))


def test_adresse_unvollstaendig_warnt() -> None:
    warnungen = describe_address(Job(id="1", title="T", company="Personalberatung"))
    assert warnungen and "Straße" in warnungen[0] and "PLZ/Ort" in warnungen[0]


def test_vollstaendige_adresse_warnt_nicht() -> None:
    job = Job(id="1", title="T", company="A", street="Beethovengasse 43", postal_city="2340 Mödling")
    assert describe_address(job) == []


def test_stellen_datei_braucht_pflichtfelder(tmp_path: Path) -> None:
    path = tmp_path / "stellen.yaml"
    path.write_text("- firma: ACME\n", encoding="utf-8")
    with pytest.raises(ConfigError, match="titel"):
        load_jobs_file(path)


def _profil_mit_suche(tmp_path: Path, suche_yaml: str) -> Path:
    pfad = tmp_path / "profil.yaml"
    pfad.write_text(
        "bewerber:\n  name: Max Muster\n"
        "vorlagen:\n  motivationsschreiben: b.docx\n  lebenslauf: l.docx\n"
        f"suche:\n{suche_yaml}",
        encoding="utf-8",
    )
    return pfad


def test_rollen_und_themen_werden_kombiniert(tmp_path: Path) -> None:
    """Werkstudentenstellen nennen das Studienfach oft nicht — daher nach Können suchen."""
    pfad = _profil_mit_suche(
        tmp_path, "  rollen: [Werkstudent, Praktikum]\n  themen: [VBA, CAD]\n  query: []\n"
    )
    assert load_settings(pfad).search.queries == [
        "Werkstudent VBA",
        "Werkstudent CAD",
        "Praktikum VBA",
        "Praktikum CAD",
    ]


def test_eigene_begriffe_stehen_vor_den_kombinationen(tmp_path: Path) -> None:
    pfad = _profil_mit_suche(
        tmp_path,
        "  query: [Werkstudent Maschinenbauingenieur]\n  rollen: [Werkstudent]\n  themen: [SQL]\n",
    )
    assert load_settings(pfad).search.queries == [
        "Werkstudent Maschinenbauingenieur",
        "Werkstudent SQL",
    ]


def test_doppelte_suchbegriffe_werden_zusammengefasst(tmp_path: Path) -> None:
    pfad = _profil_mit_suche(
        tmp_path, "  query: [Werkstudent CAD]\n  rollen: [Werkstudent]\n  themen: [CAD, CAD]\n"
    )
    assert load_settings(pfad).search.queries == ["Werkstudent CAD"]


def test_ohne_rollen_bleibt_es_bei_den_eigenen_begriffen(tmp_path: Path) -> None:
    pfad = _profil_mit_suche(tmp_path, "  query: [Werkstudent Maschinenbau]\n  themen: [CAD]\n")
    assert load_settings(pfad).search.queries == ["Werkstudent Maschinenbau"]


def test_fehlende_konfiguration_meldet_pfad(tmp_path: Path) -> None:
    with pytest.raises(ConfigError, match="profil.example.yaml"):
        load_settings(tmp_path / "fehlt.yaml")


def test_stub_composer_liefert_vollstaendiges_ergebnis(
    tmp_path: Path, cv_template: Path, letter_template: Path
) -> None:
    profil = tmp_path / "profil.yaml"
    profil.write_text(
        f"""
bewerber:
  name: Max Muster
vorlagen:
  motivationsschreiben: {letter_template}
  lebenslauf: {cv_template}
schwerpunkte:
  - key: konstruktion
    label: Konstruktion
    keywords: [3d-druck]
    hervorheben: [GitHub]
""",
        encoding="utf-8",
    )
    settings = load_settings(profil)
    job = Job(id="1", title="Aushilfe 3D-Druck", company="ACME GmbH", description="3d-druck")
    template = TemplateData(
        facts="fakten", projects=["Drucker", "GitHub"], skills=["CAD: Creo"]
    )
    focus, letter, adaptation = StubComposer().compose(settings, job, template)
    assert focus.key == "konstruktion"
    assert letter.paragraphs and "TESTMODUS" in letter.paragraphs[0]
    assert adaptation.project_order[0] == "GitHub"  # 'hervorheben' zieht nach vorne


# ------------------------------------------------------------------ Ausbildung


def test_ausbildung_bullets_werden_umformuliert(cv_template: Path, tmp_path: Path) -> None:
    adaptation = CVAdaptation(
        education_edits=[
            SectionEdit(
                anchor="2023 – dato",
                bullets=["Bachelorarbeit: Schutzeinhausung, konstruiert in Fusion 360"],
            )
        ]
    )
    output, warnings = cv_render.render_cv(cv_template, adaptation, tmp_path / "out.docx")
    assert warnings == []
    text = "\n".join(p.text for p in Document(str(output)).tables[0].rows[0].cells[1].paragraphs)
    assert "Schutzeinhausung" in text
    assert "Bachelorstudium Maschinenbau, TU Wien" in text  # Titelzeile bleibt


def test_ausbildungseintrag_wird_nie_entfernt(cv_template: Path, tmp_path: Path) -> None:
    """Ein gelöschter Ausbildungseintrag wäre eine Lücke im Lebenslauf."""
    adaptation = CVAdaptation(
        education_edits=[SectionEdit(anchor="2023 – dato", bullets=[], drop=True)]
    )
    output, warnings = cv_render.render_cv(cv_template, adaptation, tmp_path / "out.docx")
    assert any("abgelehnt" in w for w in warnings)
    assert len(cv_render.education_anchors(output)) == 1


def test_tabellenraster_folgt_den_zellbreiten(cv_template: Path, tmp_path: Path) -> None:
    """Ein Raster, das nicht zu den Zellen passt, quetscht den Text bei der PDF-Ausgabe."""
    document = Document(str(cv_template))
    grid = document.tables[0]._tbl.find(qn("w:tblGrid"))
    for column in grid.findall(qn("w:gridCol")):
        column.set(qn("w:w"), "100")
    document.save(str(tmp_path / "schief.docx"))

    output, _ = cv_render.render_cv(tmp_path / "schief.docx", CVAdaptation(), tmp_path / "out.docx")
    table = Document(str(output)).tables[0]
    breiten = [c.get(qn("w:w")) for c in table._tbl.find(qn("w:tblGrid")).findall(qn("w:gridCol"))]
    zellen = [c._tc.xpath("./w:tcPr/w:tcW/@w:w")[0] for c in table.rows[0].cells]
    assert breiten == zellen


def test_raster_ohne_zellbreiten_bleibt_unberuehrt(cv_template: Path) -> None:
    document = Document(str(cv_template))
    table = document.tables[0]
    for cell in table.rows[0].cells:
        for tcw in cell._tc.xpath("./w:tcPr/w:tcW"):
            tcw.getparent().remove(tcw)
    assert dt.sync_table_grid(table) is False


def test_kennung_aus_education_anchors_trifft(cv_template: Path, tmp_path: Path) -> None:
    """Die KI bekommt 'Kennung | Titelzeile' — genau damit muss sie treffen."""
    anchor = cv_render.education_anchors(cv_template)[0]
    adaptation = CVAdaptation(
        education_edits=[SectionEdit(anchor=anchor, bullets=["Bachelorarbeit: Schutzeinhausung"])]
    )
    output, warnings = cv_render.render_cv(cv_template, adaptation, tmp_path / "out.docx")
    assert warnings == []
    text = "\n".join(p.text for p in Document(str(output)).tables[0].rows[0].cells[1].paragraphs)
    assert "Schutzeinhausung" in text


def test_ausbildungskennungen_mit_titelzeile(cv_template: Path) -> None:
    anchors = cv_render.education_anchors(cv_template)
    assert anchors == ["2023 – dato | Bachelorstudium Maschinenbau, TU Wien"]


def test_ausbildung_ohne_anpassung_bleibt_unveraendert(
    cv_template: Path, tmp_path: Path
) -> None:
    output, warnings = cv_render.render_cv(cv_template, CVAdaptation(), tmp_path / "out.docx")
    assert warnings == []
    text = "\n".join(p.text for p in Document(str(output)).tables[0].rows[0].cells[1].paragraphs)
    assert "Bachelorarbeit zu additiver Fertigung" in text


# ---------------------------------------------------------- Antwort auswerten


def _settings_for_parse(tmp_path: Path, cv_template: Path, letter_template: Path):
    profil = tmp_path / "profil.yaml"
    profil.write_text(
        f"""
bewerber:
  name: Max Muster
vorlagen:
  motivationsschreiben: {letter_template}
  lebenslauf: {cv_template}
schwerpunkte:
  - key: it
    label: IT
    keywords: [python]
    hervorheben: [GitHub]
""",
        encoding="utf-8",
    )
    return load_settings(profil)


def test_antwort_wird_vollstaendig_uebernommen(
    tmp_path: Path, cv_template: Path, letter_template: Path
) -> None:
    settings = _settings_for_parse(tmp_path, cv_template, letter_template)
    job = Job(id="1", title="Werkstudent", company="ACME GmbH", location="1010 Wien")
    data = {
        "schwerpunkt": "it",
        "schwerpunkt_begruendung": "Python im Anforderungsprofil",
        "betreff": "Werkstudent Datenanalyse",
        "anrede": "Sehr geehrte Frau Muster,",
        "absaetze": ["Erster\nAbsatz", "  ", "Zweiter Absatz"],
        "ausbildung_bullets": [{"eintrag": "2023 – dato", "bullets": ["Bachelorarbeit zu CAD"]}],
        "projekt_reihenfolge": ["GitHub", "Drucker"],
        "projekt_bullets": [{"projekt": "GitHub", "bullets": ["Python-Projekte", ""]}],
        "kenntnis_zeilen": ["CAD: Creo", ""],
    }
    focus, letter, adaptation = parse_response(settings, job, data)

    assert focus.key == "it" and focus.emphasise == ["GitHub"]
    assert letter.paragraphs == ["Erster Absatz", "Zweiter Absatz"]  # leer raus, Umbruch geglättet
    assert letter.subject == "Werkstudent Datenanalyse"
    assert letter.company_postal_city == "1010 Wien"
    assert adaptation.education_edits[0].anchor == "2023 – dato"
    assert adaptation.project_edits[0].bullets == ["Python-Projekte"]
    assert adaptation.skill_lines == ["CAD: Creo"]


def test_antwort_ohne_absaetze_ist_ein_fehler(
    tmp_path: Path, cv_template: Path, letter_template: Path
) -> None:
    settings = _settings_for_parse(tmp_path, cv_template, letter_template)
    job = Job(id="1", title="Werkstudent", company="ACME GmbH")
    with pytest.raises(ComposerError, match="keinen Brieftext"):
        parse_response(settings, job, {"absaetze": []})


def test_unbekannter_schwerpunkt_faellt_auf_heuristik_zurueck(
    tmp_path: Path, cv_template: Path, letter_template: Path
) -> None:
    settings = _settings_for_parse(tmp_path, cv_template, letter_template)
    job = Job(id="1", title="Werkstudent", company="ACME", description="python")
    hint = detect_focus(job, settings.focus_rules)
    focus, _, _ = parse_response(
        settings, job, {"schwerpunkt": "gibt-es-nicht", "absaetze": ["Text"]}, hint
    )
    assert focus.key == "it"


def test_antworttext_aus_chunks() -> None:
    """Mistral liefert content je nach Modell als String oder als Chunk-Liste."""
    assert message_text("fertig") == "fertig"
    assert message_text([{"text": "a"}, {"text": "b"}]) == "ab"
    assert message_text(None) == ""


# ------------------------------------------------------------- KI-Vorauswahl


def _pairs(*jobs: Job) -> list[tuple[Job, int, int]]:
    """Treffer im Format von rank_jobs: (Stelle, Schwerpunkt, Titel)."""
    return [(job, 1, 2) for job in jobs]


def test_vorauswahl_uebernimmt_urteil_und_begruendung() -> None:
    batch = _pairs(
        Job(id="1", title="Werkstudent CAD", company="A"),
        Job(id="2", title="Senior Manager", company="B"),
    )
    data = {
        "bewertungen": [
            {"nr": 2, "passend": False, "punkte": 10, "grund": "x", "begruendung": "Vollzeit"},
            {"nr": 1, "passend": True, "punkte": 85, "begruendung": "CAD  passt\nzum Profil"},
        ]
    }
    urteile = screening.parse_response(data, batch, min_score=60)
    assert [u.job.id for u in urteile] == ["1", "2"]  # Reihenfolge des Stapels
    assert urteile[0].fits and urteile[0].score == 85
    assert urteile[0].reason == "CAD passt zum Profil"  # Umbruch geglättet
    assert urteile[0].title_hits == 2 and urteile[0].focus_score == 1
    assert not urteile[1].fits


def test_mindestpunkte_ueberstimmen_ein_zu_gutmuetiges_ja() -> None:
    batch = _pairs(Job(id="1", title="Werkstudent", company="A"))
    data = {"bewertungen": [{"nr": 1, "passend": True, "punkte": 45, "begruendung": "Randfall"}]}
    assert not screening.parse_response(data, batch, min_score=60)[0].fits
    assert screening.parse_response(data, batch, min_score=40)[0].fits


def test_uebergangene_stelle_faellt_nicht_unter_den_tisch() -> None:
    """Bewertet das Modell eine Anzeige nicht, muss das auffallen."""
    batch = _pairs(
        Job(id="1", title="Werkstudent", company="A"),
        Job(id="2", title="Praktikum", company="B"),
    )
    urteile = screening.parse_response(
        {"bewertungen": [{"nr": 1, "passend": True, "punkte": 90, "begruendung": "passt"}]},
        batch,
    )
    assert len(urteile) == 2
    assert urteile[1].rated is False and urteile[1].fits is False
    assert "übergangen" in urteile[1].reason


def test_erfundene_nummern_werden_ignoriert() -> None:
    batch = _pairs(Job(id="1", title="Werkstudent", company="A"))
    urteile = screening.parse_response(
        {"bewertungen": [{"nr": 7, "passend": True, "punkte": 90, "begruendung": "?"}]}, batch
    )
    assert urteile[0].rated is False


def test_punkte_bleiben_in_der_spanne() -> None:
    batch = _pairs(Job(id="1", title="W", company="A"), Job(id="2", title="P", company="B"))
    urteile = screening.parse_response(
        {
            "bewertungen": [
                {"nr": 1, "passend": True, "punkte": 250, "begruendung": ""},
                {"nr": 2, "passend": False, "punkte": "keine Zahl", "begruendung": ""},
            ]
        },
        batch,
    )
    assert urteile[0].score == 100 and urteile[1].score == 0


def test_stapel_statt_einer_anfrage_je_anzeige() -> None:
    """Der Sinn der Vorauswahl: viele Anzeigen, wenige Anfragen."""
    jobs = _pairs(*(Job(id=str(i), title=f"Stelle {i}", company="A") for i in range(70)))
    stapel = screening.batches(jobs)
    assert [len(s) for s in stapel] == [30, 30, 10]
    assert screening.anfragen(len(jobs)) == 3
    assert screening.anfragen(0) == 0


def test_prompt_nummeriert_und_kuerzt_lange_ausschreibungen(
    tmp_path: Path, cv_template: Path, letter_template: Path
) -> None:
    settings = _settings_for_parse(tmp_path, cv_template, letter_template)
    batch = _pairs(
        Job(id="1", title="Werkstudent CAD", company="ACME", description="wort " * 900),
        Job(id="2", title="Praktikum Python", company="Beta", location="Wien"),
    )
    prompt = screening.build_prompt(settings, TemplateData(facts="Faktenblatt"), batch)
    assert "Faktenblatt" in prompt
    assert "1. Werkstudent CAD — ACME" in prompt and "2. Praktikum Python — Beta" in prompt
    assert "[…]" in prompt  # der lange Text wurde gekürzt
    assert len(prompt) < 3000


def test_stub_screener_sortiert_nichts_aus(
    tmp_path: Path, cv_template: Path, letter_template: Path
) -> None:
    settings = _settings_for_parse(tmp_path, cv_template, letter_template)
    pairs = _pairs(Job(id="1", title="Werkstudent", company="A"))
    urteile = screening.StubScreener().screen(settings, pairs, TemplateData())
    assert urteile[0].fits and urteile[0].rated is False


def test_suchlauf_legt_beide_ergebnisse_ab(
    tmp_path: Path, cv_template: Path, letter_template: Path
) -> None:
    settings = _settings_for_parse(tmp_path, cv_template, letter_template)
    job_a = Job(id="1", title="Werkstudent CAD", company="ACME", url="https://example.com/a")
    job_b = Job(id="2", title="Senior Manager", company="Beta")
    pairs = _pairs(job_a, job_b)
    urteile = [
        Screening(job=job_a, fits=True, score=85, reason="passt", title_hits=2, focus_score=1),
        Screening(job=job_b, fits=False, score=10, reason="Vollzeit"),
    ]
    ordner = save_search_run(settings, pairs, urteile, ["Careerjet down"], model="testmodell",
                             root=tmp_path / "suchlaeufe")

    suche = json.loads((ordner / "suche.json").read_text(encoding="utf-8"))
    assert suche["anzahl"] == 2 and suche["probleme"] == ["Careerjet down"]
    assert suche["stellen"][0]["stelle"]["title"] == "Werkstudent CAD"
    assert suche["stellen"][0]["platz"] == 1

    vorauswahl = json.loads((ordner / "vorauswahl.json").read_text(encoding="utf-8"))
    assert vorauswahl["modell"] == "testmodell" and vorauswahl["passend_anzahl"] == 1
    assert vorauswahl["passend"][0]["stelle"]["url"] == "https://example.com/a"
    assert vorauswahl["aussortiert"][0]["firma"] == "Beta"


def test_suchlauf_ohne_vorauswahl_schreibt_nur_die_suche(
    tmp_path: Path, cv_template: Path, letter_template: Path
) -> None:
    settings = _settings_for_parse(tmp_path, cv_template, letter_template)
    ordner = save_search_run(
        settings, _pairs(Job(id="1", title="W", company="A")), None, [],
        root=tmp_path / "suchlaeufe",
    )
    assert (ordner / "suche.json").is_file()
    assert not (ordner / "vorauswahl.json").exists()


def test_zwei_laeufe_in_derselben_sekunde_vermischen_sich_nicht(
    tmp_path: Path, cv_template: Path, letter_template: Path
) -> None:
    settings = _settings_for_parse(tmp_path, cv_template, letter_template)
    pairs = _pairs(Job(id="1", title="W", company="A"))
    wurzel = tmp_path / "suchlaeufe"
    erster = save_search_run(settings, pairs, None, [], root=wurzel)
    zweiter = save_search_run(settings, pairs, [], [], root=wurzel)
    assert erster != zweiter
    assert not (erster / "vorauswahl.json").exists()  # der zweite Lauf hat nichts überschrieben
    assert (zweiter / "vorauswahl.json").is_file()
