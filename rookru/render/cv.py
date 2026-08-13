"""Lebenslauf: Vorlage übernehmen und nur die freigegebenen Abschnitte anpassen.

Angepasst werden AUSBILDUNG (nur Stichpunkte), PROJEKTE (Reihenfolge und
Stichpunkte) und BESONDERE KENNTNISSE UND FÄHIGKEITEN. Persönliche Daten,
Berufserfahrungen und Sprachkenntnisse bleiben Zeichen für Zeichen so, wie sie
in der Vorlage stehen.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.document import Document as DocumentType
from docx.table import Table

from ..models import CVAdaptation
from . import docx_tools as dt

HEADINGS = [
    "PERSÖNLICHE DATEN",
    "AUSBILDUNG",
    "BERUFSERFAHRUNGEN",
    "PROJEKTE",
    "SPRACHKENNTNISSE",
    "BESONDERE KENNTNISSE UND FÄHIGKEITEN",
]
EDUCATION = "AUSBILDUNG"
PROJECTS = "PROJEKTE"
SKILLS = "BESONDERE KENNTNISSE UND FÄHIGKEITEN"

EDITABLE_SECTIONS = (EDUCATION, PROJECTS, SKILLS)


def _entry(table: Table) -> tuple[str, object] | None:
    """Kennung und Inhaltszelle einer Eintragstabelle.

    Gibt None zurück, wenn die Tabelle nicht dem erwarteten Aufbau entspricht
    (mindestens eine Zeile mit zwei Spalten). Fremde Tabellen — etwa ein
    Layoutrahmen im Kopf des Dokuments — werden so übersprungen statt den
    ganzen Lauf abzubrechen.
    """
    if not table.rows or not table.columns:
        return None
    cells = table.rows[0].cells
    if len(cells) < 2:
        return None
    return cells[0].text.strip(), cells[1]


def _match_table(tables: list[Table], anchor: str) -> Table | None:
    needle = anchor.strip().lower()
    if not needle:
        return None
    for table in tables:
        entry = _entry(table)
        if entry and entry[0].lower() == needle:
            return table
    for table in tables:
        entry = _entry(table)
        if entry and needle in f"{entry[0]} {entry[1].text}".lower():
            return table
    return None


def _section_tables(document: DocumentType, section: str) -> list[Table]:
    return dt.section_tables(document, section, [h for h in HEADINGS if h != section])


def apply_education(document: DocumentType, adaptation: CVAdaptation) -> list[str]:
    """Stichpunkte der Ausbildungseinträge umformulieren.

    Reihenfolge und Bestand bleiben unangetastet: Ein Lebenslauf ohne
    chronologische Ausbildung oder mit fehlendem Abschluss fällt auf.
    """
    if not adaptation.education_edits:
        return []

    warnings: list[str] = []
    tables = _section_tables(document, EDUCATION)
    if not tables:
        return ["Abschnitt AUSBILDUNG in der Vorlage nicht gefunden — unverändert übernommen."]

    for edit in adaptation.education_edits:
        table = _match_table(tables, edit.anchor)
        if table is None:
            warnings.append(
                f"Ausbildungseintrag '{edit.anchor}' nicht in der Vorlage gefunden — übersprungen."
            )
            continue
        if edit.drop:
            warnings.append(
                f"Ausbildungseintrag '{edit.anchor}' sollte entfernt werden — abgelehnt, "
                "Ausbildungseinträge bleiben vollständig."
            )
            continue
        entry = _entry(table)
        if edit.bullets and entry:
            dt.set_bullets(entry[1], edit.bullets)
    return warnings


def apply_projects(document: DocumentType, adaptation: CVAdaptation) -> list[str]:
    """Projekte umsortieren und Stichpunkte ersetzen. Gibt Warnungen zurück."""
    warnings: list[str] = []
    tables = _section_tables(document, PROJECTS)
    if not tables:
        return ["Abschnitt PROJEKTE in der Vorlage nicht gefunden — unverändert übernommen."]

    for edit in adaptation.project_edits:
        table = _match_table(tables, edit.anchor)
        if table is None:
            warnings.append(f"Projekt '{edit.anchor}' nicht in der Vorlage gefunden — übersprungen.")
            continue
        if edit.drop:
            dt.delete_table(table)
            tables = [t for t in tables if t is not table]
            continue
        entry = _entry(table)
        if edit.bullets and entry:
            dt.set_bullets(entry[1], edit.bullets)

    if adaptation.project_order:
        ordered: list[Table] = []
        for anchor in adaptation.project_order:
            table = _match_table(tables, anchor)
            if table is not None and table not in ordered:
                ordered.append(table)
        ordered += [t for t in tables if t not in ordered]
        if len(ordered) == len(tables):
            dt.reorder_tables(tables, ordered)
        else:
            warnings.append("Projektreihenfolge konnte nicht angewendet werden.")

    return warnings


def apply_skills(document: DocumentType, adaptation: CVAdaptation) -> list[str]:
    """Zeilen unter BESONDERE KENNTNISSE ersetzen."""
    if not adaptation.skill_lines:
        return []

    blocks = dt.section_blocks(document, SKILLS, [])
    paragraphs = [b for b in blocks if hasattr(b, "runs") and b.text.strip()]
    if not paragraphs:
        return ["Abschnitt BESONDERE KENNTNISSE nicht gefunden — unverändert übernommen."]

    lines = adaptation.skill_lines
    for paragraph, text in zip(paragraphs, lines):
        dt.set_text_preserving(paragraph, text)

    if len(lines) > len(paragraphs):
        anchor = paragraphs[-1]
        for text in lines[len(paragraphs) :]:
            anchor = dt.clone_paragraph_after(anchor, text)

    for paragraph in paragraphs[len(lines) :]:
        dt.delete_paragraph(paragraph)

    return []


def render_cv(template: Path, adaptation: CVAdaptation, output: Path) -> tuple[Path, list[str]]:
    """Erzeugt den angepassten Lebenslauf; gibt Pfad und Warnungen zurück."""
    document = Document(str(template))
    warnings = apply_education(document, adaptation)
    warnings += apply_projects(document, adaptation)
    warnings += apply_skills(document, adaptation)

    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output))
    return output, warnings


def read_facts(document: DocumentType) -> tuple[str, list[str]]:
    """Faktenblatt für die KI plus Hinweise auf nicht lesbare Tabellen."""
    lines: list[str] = []
    issues: list[str] = []
    tabelle_nr = 0

    for block in dt.iter_block_items(document):
        if isinstance(block, Table):
            tabelle_nr += 1
            entry = _entry(block)
            if entry is None:
                text = " ".join(block._tbl.itertext()).split()
                vorschau = " ".join(text)[:60] or "leer"
                spalten = len(block.columns) if block.columns else 0
                issues.append(
                    f"Tabelle {tabelle_nr} hat {spalten} Spalte(n) statt zwei "
                    f"und wurde übersprungen (Inhalt: „{vorschau}“)."
                )
                continue
            left, cell = entry
            paragraphs = [p.text.strip() for p in cell.paragraphs if p.text.strip()]
            if paragraphs:
                lines.append(f"- [{left}] {paragraphs[0]}")
                lines.extend(f"    {p}" for p in paragraphs[1:])
        else:
            text = block.text.strip()
            if not text:
                continue
            if text.upper() in HEADINGS:
                lines.append(f"\n## {text.upper()}")
            else:
                lines.append(text)

    return "\n".join(lines).strip(), issues


def template_facts(template: Path) -> str:
    """Liest die Vorlage als Faktenblatt für die KI aus.

    Damit gibt es nur eine Quelle der Wahrheit: den Lebenslauf selbst.
    """
    return read_facts(Document(str(template)))[0]


def project_anchors(template: Path) -> list[str]:
    """Kennungen der Projekt-Zeilen (linke Tabellenspalte) in Vorlagenreihenfolge."""
    document = Document(str(template))
    entries = (_entry(t) for t in _section_tables(document, PROJECTS))
    return [e[0] for e in entries if e]


def education_anchors(template: Path) -> list[str]:
    """Kennungen der Ausbildungseinträge, jeweils mit Titelzeile zur Orientierung."""
    document = Document(str(template))
    anchors = []
    for table in _section_tables(document, EDUCATION):
        entry = _entry(table)
        if entry is None:
            continue
        left, cell = entry
        head = next((p.text.strip() for p in cell.paragraphs if p.text.strip()), "")
        anchors.append(f"{left} | {head}")
    return anchors


def read_structure(template: Path) -> tuple[str, list[str], list[str], list[str], list[str]]:
    """Liest die Vorlage einmal komplett.

    Zurück kommen Faktenblatt, Projektkennungen, Ausbildungskennungen,
    Kenntnis-Zeilen und Hinweise auf alles, was nicht gelesen werden konnte.
    """
    document = Document(str(template))
    facts, issues = read_facts(document)

    projects = [e[0] for e in (_entry(t) for t in _section_tables(document, PROJECTS)) if e]
    if not projects:
        issues.append(f"Kein Eintrag unter {PROJECTS} gefunden — Abschnitt bleibt unverändert.")

    education = []
    for table in _section_tables(document, EDUCATION):
        entry = _entry(table)
        if entry is None:
            continue
        head = next((p.text.strip() for p in entry[1].paragraphs if p.text.strip()), "")
        education.append(f"{entry[0]} | {head}")
    if not education:
        issues.append(f"Kein Eintrag unter {EDUCATION} gefunden — Abschnitt bleibt unverändert.")

    blocks = dt.section_blocks(document, SKILLS, [])
    skills = [b.text.strip() for b in blocks if hasattr(b, "runs") and b.text.strip()]
    if not skills:
        issues.append(f"Keine Zeile unter {SKILLS} gefunden — Abschnitt bleibt unverändert.")

    return facts, projects, education, skills, issues


def skill_lines(template: Path) -> list[str]:
    document = Document(str(template))
    blocks = dt.section_blocks(document, SKILLS, [])
    return [b.text.strip() for b in blocks if hasattr(b, "runs") and b.text.strip()]
