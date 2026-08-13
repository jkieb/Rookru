"""Motivationsschreiben: Vorlage mit {{PLATZHALTERN}} befüllen."""

from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document

from ..models import LetterContent
from . import docx_tools as dt

MONATE = (
    "Januar", "Februar", "März", "April", "Mai", "Juni",
    "Juli", "August", "September", "Oktober", "November", "Dezember",
)

MAX_PARAGRAPH_SLOTS = 8


def format_date(day: date) -> str:
    """Datum in der Form '13. August 2026' — wie im Muster."""
    return f"{day.day}. {MONATE[day.month - 1]} {day.year}"


def build_mapping(content: LetterContent) -> dict[str, str]:
    mapping = {
        "{{DATUM}}": format_date(content.letter_date),
        "{{FIRMA_NAME}}": content.company_name,
        "{{FIRMA_ABTEILUNG}}": content.company_department,
        "{{FIRMA_STRASSE}}": content.company_street,
        "{{FIRMA_PLZ_ORT}}": content.company_postal_city,
        "{{STELLE}}": content.subject,
        "{{ANREDE}}": content.salutation,
    }
    for i in range(1, MAX_PARAGRAPH_SLOTS + 1):
        text = content.paragraphs[i - 1] if i <= len(content.paragraphs) else ""
        mapping[f"{{{{ABSATZ_{i}}}}}"] = text
    return mapping


def render_letter(template: Path, content: LetterContent, output: Path) -> Path:
    """Füllt die Vorlage und speichert das Ergebnis unter `output`."""
    document = Document(str(template))
    mapping = build_mapping(content)

    # Mehr Absätze als Platzhalter: den letzten belegten Absatz klonen.
    slots = [
        p
        for p in dt.all_paragraphs(document)
        if p.text.strip().startswith("{{ABSATZ_") and p.text.strip().endswith("}}")
    ]
    if len(content.paragraphs) > len(slots) and slots:
        anchor = slots[-1]
        for i in range(len(slots), len(content.paragraphs)):
            key = f"{{{{ABSATZ_{i + 1}}}}}"
            anchor = dt.clone_paragraph_after(anchor, key)
            mapping[key] = content.paragraphs[i]

    # Zeilen, für die es keinen Inhalt gibt (z. B. unbekannte Abteilung),
    # ganz entfernen statt leer stehen zu lassen.
    for paragraph in list(dt.all_paragraphs(document)):
        text = paragraph.text.strip()
        if text in mapping and not mapping[text].strip():
            dt.delete_paragraph(paragraph)

    dt.replace_placeholders(document, mapping)

    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output))
    return output


def check_placeholders(path: Path) -> set[str]:
    """Übrig gebliebene Platzhalter im erzeugten Dokument."""
    return dt.remaining_placeholders(Document(str(path)))
