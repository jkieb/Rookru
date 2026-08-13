"""Werkzeuge, um bestehende docx-Vorlagen zu befüllen und umzubauen.

Grundsatz: Die Vorlage bestimmt das Layout. Hier wird nur Text ersetzt,
umsortiert oder entfernt — Schriftart, Größen, Einzüge und Tabellenbreiten
bleiben die der Vorlage.
"""

from __future__ import annotations

import copy
import re
from typing import Callable, Iterator

from docx.document import Document as DocumentType
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

BULLET_PREFIX = "●"
PLACEHOLDER_RE = re.compile(r"\{\{[A-Z0-9_]+\}\}")


def iter_block_items(parent) -> Iterator[Paragraph | Table]:
    """Absätze und Tabellen eines Dokuments oder einer Zelle in Reihenfolge."""
    if isinstance(parent, _Cell):
        element = parent._tc
    else:
        element = parent.element.body
    for child in element.iterchildren():
        tag = child.tag.split("}")[-1]
        if tag == "p":
            yield Paragraph(child, parent)
        elif tag == "tbl":
            yield Table(child, parent)


def all_paragraphs(document: DocumentType) -> Iterator[Paragraph]:
    """Alle Absätze, auch die in Tabellenzellen."""
    for block in iter_block_items(document):
        if isinstance(block, Paragraph):
            yield block
        else:
            for row in block.rows:
                for cell in row.cells:
                    yield from cell.paragraphs


def set_text_preserving(paragraph: Paragraph, text: str) -> None:
    """Ersetzt den Absatztext und behält die Formatierung des ersten Runs."""
    runs = paragraph.runs
    if not runs:
        paragraph.add_run(text)
        return
    runs[0].text = text
    for run in runs[1:]:
        run.text = ""


def replace_in_paragraph(paragraph: Paragraph, mapping: dict[str, str]) -> bool:
    """Ersetzt Platzhalter, auch wenn Word sie über mehrere Runs verteilt hat.

    Der Ersatztext übernimmt die Formatierung des Runs, in dem der Platzhalter
    beginnt; der Rest des Platzhalters wird aus den Folge-Runs gelöscht.
    """
    runs = paragraph.runs
    if not runs:
        return False
    full = "".join(run.text for run in runs)
    if not any(key in full for key in mapping):
        return False

    # Zeichenposition → Run-Index
    spans: list[tuple[int, int]] = []
    pos = 0
    for run in runs:
        spans.append((pos, pos + len(run.text)))
        pos += len(run.text)

    changed = False
    for key, value in mapping.items():
        while True:
            full = "".join(run.text for run in runs)
            start = full.find(key)
            if start < 0:
                break
            end = start + len(key)
            texts = [run.text for run in runs]
            offset = 0
            for i, run in enumerate(runs):
                r_start, r_end = offset, offset + len(texts[i])
                offset = r_end
                if r_end <= start or r_start >= end:
                    continue
                local_start = max(start - r_start, 0)
                local_end = min(end - r_start, len(texts[i]))
                head = texts[i][:local_start]
                tail = texts[i][local_end:]
                if r_start <= start < r_end:
                    run.text = head + value + tail
                else:
                    run.text = head + tail
            changed = True
    return changed


def replace_placeholders(document: DocumentType, mapping: dict[str, str]) -> set[str]:
    """Ersetzt alle Platzhalter im Dokument; gibt die ersetzten Schlüssel zurück."""
    used: set[str] = set()
    for paragraph in all_paragraphs(document):
        before = paragraph.text
        if replace_in_paragraph(paragraph, mapping):
            used.update(key for key in mapping if key in before)
    return used


def remaining_placeholders(document: DocumentType) -> set[str]:
    """Platzhalter, die nach dem Befüllen noch im Dokument stehen."""
    found: set[str] = set()
    for paragraph in all_paragraphs(document):
        found.update(PLACEHOLDER_RE.findall(paragraph.text))
    return found


def delete_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._p
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def delete_table(table: Table, also_spacer: bool = True) -> None:
    """Entfernt eine Tabelle und den zugehörigen Leerabsatz."""
    element = table._tbl
    parent = element.getparent()
    if parent is None:
        return
    if also_spacer:
        for neighbour in (element.getnext(), element.getprevious()):
            if neighbour is not None and neighbour.tag == qn("w:p"):
                if not "".join(neighbour.itertext()).strip():
                    parent.remove(neighbour)
                    break
    parent.remove(element)


def clone_paragraph_after(paragraph: Paragraph, text: str) -> Paragraph:
    """Kopiert einen Absatz (inkl. Format) und setzt neuen Text hinein."""
    new_element = copy.deepcopy(paragraph._p)
    paragraph._p.addnext(new_element)
    clone = Paragraph(new_element, paragraph._parent)
    set_text_preserving(clone, text)
    return clone


def find_paragraph(document: DocumentType, predicate: Callable[[str], bool]) -> Paragraph | None:
    for paragraph in iter_block_items(document):
        if isinstance(paragraph, Paragraph) and predicate(paragraph.text.strip()):
            return paragraph
    return None


def section_blocks(
    document: DocumentType, heading: str, stop_headings: list[str]
) -> list[Paragraph | Table]:
    """Alle Blöcke zwischen einer Überschrift und der nächsten Überschrift."""
    blocks: list[Paragraph | Table] = []
    collecting = False
    stops = {s.strip().upper() for s in stop_headings}
    for block in iter_block_items(document):
        if isinstance(block, Paragraph):
            text = block.text.strip().upper()
            if text == heading.strip().upper():
                collecting = True
                continue
            if collecting and text in stops:
                break
        if collecting:
            blocks.append(block)
    return blocks


def section_tables(
    document: DocumentType, heading: str, stop_headings: list[str]
) -> list[Table]:
    return [b for b in section_blocks(document, heading, stop_headings) if isinstance(b, Table)]


def table_anchor(table: Table) -> str:
    """Text der linken Spalte — dient als Kennung einer Vorlagenzeile."""
    return table.rows[0].cells[0].text.strip()


def sync_table_grid(table: Table) -> bool:
    """Gleicht w:tblGrid an die Spaltenbreiten der ersten Zeile an.

    Word rechnet sich die Spalten aus den Zellbreiten zusammen und pflegt das
    Raster oft gar nicht; LibreOffice nimmt beim Umbruch nach PDF aber das
    Raster ernst. Steht dort etwas anderes als in den Zellen — etwa zwei
    gleich breite Spalten statt schmal/breit — rutscht der Text in eine viel
    zu enge Spalte. Angeglichen wird nur, was die Vorlage selbst vorgibt.
    """
    grid = table._tbl.find(qn("w:tblGrid"))
    if grid is None:
        return False
    columns = grid.findall(qn("w:gridCol"))
    cells = table.rows[0].cells if table.rows else []
    if not columns or len(columns) != len(cells):
        return False

    widths: list[str] = []
    for cell in cells:
        if cell._tc.xpath("./w:tcPr/w:gridSpan"):  # verbundene Zellen: Finger weg
            return False
        value = cell._tc.xpath("./w:tcPr/w:tcW/@w:w")
        unit = cell._tc.xpath("./w:tcPr/w:tcW/@w:type")
        if not value or not unit or unit[0] != "dxa" or not value[0].isdigit():
            return False
        widths.append(value[0])

    if [c.get(qn("w:w")) for c in columns] == widths:
        return False
    for column, width in zip(columns, widths):
        column.set(qn("w:w"), width)
    return True


def sync_table_grids(document: DocumentType) -> int:
    """Gleicht alle Tabellenraster an; gibt die Zahl der Korrekturen zurück."""
    return sum(sync_table_grid(table) for table in document.tables)


def reorder_tables(tables: list[Table], order: list[Table]) -> None:
    """Sortiert Tabellen um und lässt die Leerabsätze dazwischen unangetastet."""
    if len(tables) != len(order) or not tables:
        return
    body = tables[0]._tbl.getparent()
    if body is None:
        return

    markers = []
    for table in tables:
        marker = OxmlElement("w:p")  # Platzhalter, der die Position freihält
        table._tbl.addprevious(marker)
        markers.append(marker)
    for table in tables:
        body.remove(table._tbl)
    for marker, table in zip(markers, order):
        marker.addprevious(table._tbl)
    for marker in markers:
        body.remove(marker)


def cell_bullets(cell: _Cell) -> list[Paragraph]:
    return [p for p in cell.paragraphs if p.text.strip().startswith(BULLET_PREFIX)]


def set_bullets(cell: _Cell, bullets: list[str]) -> None:
    """Ersetzt die Stichpunkte einer Tabellenzelle.

    Vorhandene Stichpunkt-Absätze werden wiederverwendet (Format bleibt),
    überzählige entfernt, fehlende aus dem letzten geklont.
    """
    existing = cell_bullets(cell)
    if not existing:
        return
    prefix = existing[0].text[: len(BULLET_PREFIX)] + "  "

    for paragraph, text in zip(existing, bullets):
        set_text_preserving(paragraph, prefix + text.lstrip("•●- ").strip())

    if len(bullets) > len(existing):
        anchor = existing[-1]
        for text in bullets[len(existing) :]:
            anchor = clone_paragraph_after(anchor, prefix + text.lstrip("•●- ").strip())

    for paragraph in existing[len(bullets) :]:
        delete_paragraph(paragraph)


def _scale_doc_defaults(
    document: DocumentType, factor: float, spacing_factor: float, min_pt: float
) -> None:
    """Skaliert die Grundeinstellungen aus styles.xml.

    Viele Vorlagen — auch die hier verwendeten — setzen Schriftgröße und
    Absatzabstände nicht je Absatz, sondern nur in den docDefaults. Ohne diesen
    Schritt bliebe die Verkleinerung in Tabellenzellen wirkungslos.
    """
    defaults = document.styles.element.find(qn("w:docDefaults"))
    if defaults is None:
        return

    p_default = defaults.find(qn("w:pPrDefault"))
    if p_default is not None:
        p_pr = p_default.find(qn("w:pPr"))
        spacing = p_pr.find(qn("w:spacing")) if p_pr is not None else None
        if spacing is not None:
            for attr in ("w:before", "w:after"):
                value = spacing.get(qn(attr))
                if value and value.lstrip("-").isdigit():
                    spacing.set(qn(attr), str(int(int(value) * spacing_factor)))
            line = spacing.get(qn("w:line"))
            if line and line.isdigit() and spacing.get(qn("w:lineRule")) != "exact":
                spacing.set(qn("w:line"), str(max(int(int(line) * spacing_factor), 200)))

    r_default = defaults.find(qn("w:rPrDefault"))
    if r_default is not None:
        r_pr = r_default.find(qn("w:rPr"))
        if r_pr is not None:
            for tag in ("w:sz", "w:szCs"):
                size = r_pr.find(qn(tag))
                if size is not None:
                    value = size.get(qn("w:val"))
                    if value and value.isdigit():
                        # Halbe Punkte
                        size.set(qn("w:val"), str(max(int(int(value) * factor), int(min_pt * 2))))


def scale_document(
    document: DocumentType,
    factor: float,
    spacing_factor: float | None = None,
    min_pt: float = 7.0,
) -> None:
    """Verkleinert Schrift und Abstände, um ein Dokument auf eine Seite zu ziehen.

    `spacing_factor` trennt Abstände von der Schriftgröße: So lässt sich zuerst
    nur die Luft zwischen den Absätzen wegnehmen, bevor die Schrift kleiner wird
    — das hält das Schriftbild der Vorlage länger unverändert.
    """
    spacing_factor = factor if spacing_factor is None else spacing_factor
    if factor >= 1.0 and spacing_factor >= 1.0:
        return

    _scale_doc_defaults(document, factor, spacing_factor, min_pt)

    for style_name in ("Normal", "Standard"):
        try:
            style = document.styles[style_name]
        except KeyError:
            continue
        if style.font.size is not None:
            style.font.size = Pt(max(style.font.size.pt * factor, min_pt))

    for paragraph in all_paragraphs(document):
        fmt = paragraph.paragraph_format
        if fmt.space_before is not None:
            fmt.space_before = Pt(fmt.space_before.pt * spacing_factor)
        if fmt.space_after is not None:
            fmt.space_after = Pt(fmt.space_after.pt * spacing_factor)
        for run in paragraph.runs:
            if run.font.size is not None:
                run.font.size = Pt(max(run.font.size.pt * factor, min_pt))

    margin_factor = max(min(factor, spacing_factor), 0.75)
    for section in document.sections:
        section.top_margin = int(section.top_margin * margin_factor)
        section.bottom_margin = int(section.bottom_margin * margin_factor)
        section.left_margin = int(section.left_margin * margin_factor)
        section.right_margin = int(section.right_margin * margin_factor)
